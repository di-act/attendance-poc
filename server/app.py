from flask import Flask, request, jsonify, send_file
from werkzeug.utils import secure_filename
import os
import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from datetime import datetime
import io

app = Flask(__name__)

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
ALLOWED_EXTENSIONS = {'docx', 'csv'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max file size

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Create necessary folders
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def allowed_file(filename):
    """Check if the file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_docx(docx_path):
    """Extract text content from DOCX file"""
    doc = Document(docx_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            full_text.append(' | '.join(row_data))
    
    return full_text


def read_csv_data(csv_path):
    """Read CSV file and return DataFrame"""
    try:
        df = pd.read_csv(csv_path)
        return df
    except Exception as e:
        raise Exception(f"Error reading CSV file: {str(e)}")


def generate_xlsx(docx_content, csv_df, output_path):
    """Generate XLSX file combining DOCX and CSV data"""
    wb = Workbook()
    
    # Sheet 1: DOCX Content
    ws1 = wb.active
    ws1.title = "Document Content"
    
    # Header styling
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=12)
    
    # Write DOCX content
    ws1['A1'] = "DOCX Content"
    ws1['A1'].font = header_font
    ws1['A1'].fill = header_fill
    ws1['A1'].alignment = Alignment(horizontal='center', vertical='center')
    
    row = 2
    for line in docx_content:
        ws1[f'A{row}'] = line
        row += 1
    
    # Adjust column width
    ws1.column_dimensions['A'].width = 100
    
    # Sheet 2: CSV Data
    ws2 = wb.create_sheet(title="CSV Data")
    
    # Write CSV headers
    for col_idx, column in enumerate(csv_df.columns, start=1):
        cell = ws2.cell(row=1, column=col_idx, value=column)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Write CSV data
    for row_idx, row in enumerate(csv_df.itertuples(index=False), start=2):
        for col_idx, value in enumerate(row, start=1):
            ws2.cell(row=row_idx, column=col_idx, value=value)
    
    # Adjust column widths
    for column in ws2.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws2.column_dimensions[column_letter].width = adjusted_width
    
    # Sheet 3: Summary
    ws3 = wb.create_sheet(title="Summary")
    
    ws3['A1'] = "File Processing Summary"
    ws3['A1'].font = Font(bold=True, size=14)
    
    ws3['A3'] = "Processing Date:"
    ws3['B3'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    ws3['A4'] = "DOCX Lines:"
    ws3['B4'] = len(docx_content)
    
    ws3['A5'] = "CSV Rows:"
    ws3['B5'] = len(csv_df)
    
    ws3['A6'] = "CSV Columns:"
    ws3['B6'] = len(csv_df.columns)
    
    ws3.column_dimensions['A'].width = 20
    ws3.column_dimensions['B'].width = 30
    
    # Save workbook
    wb.save(output_path)


@app.route('/')
def index():
    """Health check endpoint"""
    return jsonify({
        'status': 'running',
        'message': 'File Upload and XLSX Generation API',
        'version': '1.0.0'
    })


@app.route('/api/upload', methods=['POST'])
def upload_files():
    """
    Upload DOCX and CSV files, process them, and return XLSX file
    
    Expected form data:
    - docx_file: DOCX file
    - csv_file: CSV file
    """
    try:
        # Check if files are present in request
        if 'docx_file' not in request.files or 'csv_file' not in request.files:
            return jsonify({
                'error': 'Both docx_file and csv_file are required'
            }), 400
        
        docx_file = request.files['docx_file']
        csv_file = request.files['csv_file']
        
        # Validate files
        if docx_file.filename == '' or csv_file.filename == '':
            return jsonify({
                'error': 'No file selected'
            }), 400
        
        if not allowed_file(docx_file.filename) or not allowed_file(csv_file.filename):
            return jsonify({
                'error': 'Invalid file type. Only DOCX and CSV files are allowed'
            }), 400
        
        # Validate file extensions match expected types
        if not docx_file.filename.lower().endswith('.docx'):
            return jsonify({
                'error': 'docx_file must be a DOCX file'
            }), 400
        
        if not csv_file.filename.lower().endswith('.csv'):
            return jsonify({
                'error': 'csv_file must be a CSV file'
            }), 400
        
        # Save uploaded files
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = secure_filename(f"{timestamp}_{docx_file.filename}")
        csv_filename = secure_filename(f"{timestamp}_{csv_file.filename}")
        
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
        csv_path = os.path.join(app.config['UPLOAD_FOLDER'], csv_filename)
        
        docx_file.save(docx_path)
        csv_file.save(csv_path)
        
        # Process files
        docx_content = extract_text_from_docx(docx_path)
        csv_df = read_csv_data(csv_path)
        
        # Generate XLSX
        output_filename = f"output_{timestamp}.xlsx"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        generate_xlsx(docx_content, csv_df, output_path)
        
        # Clean up uploaded files (optional)
        os.remove(docx_path)
        os.remove(csv_path)
        
        # Return XLSX file
        return send_file(
            output_path,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=output_filename
        )
    
    except Exception as e:
        return jsonify({
            'error': f'Error processing files: {str(e)}'
        }), 500


@app.route('/api/upload-stream', methods=['POST'])
def upload_files_stream():
    """
    Upload DOCX and CSV files, process them, and return XLSX file as stream
    (without saving to disk)
    
    Expected form data:
    - docx_file: DOCX file
    - csv_file: CSV file
    """
    try:
        # Check if files are present in request
        if 'docx_file' not in request.files or 'csv_file' not in request.files:
            return jsonify({
                'error': 'Both docx_file and csv_file are required'
            }), 400
        
        docx_file = request.files['docx_file']
        csv_file = request.files['csv_file']
        
        # Validate files
        if docx_file.filename == '' or csv_file.filename == '':
            return jsonify({
                'error': 'No file selected'
            }), 400
        
        # Process DOCX from memory
        docx_stream = io.BytesIO(docx_file.read())
        doc = Document(docx_stream)
        docx_content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                docx_content.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                docx_content.append(' | '.join(row_data))
        
        # Process CSV from memory
        csv_stream = io.BytesIO(csv_file.read())
        csv_df = pd.read_csv(csv_stream)
        
        # Generate XLSX in memory
        output_stream = io.BytesIO()
        wb = Workbook()
        
        # Sheet 1: DOCX Content
        ws1 = wb.active
        ws1.title = "Document Content"
        
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=12)
        
        ws1['A1'] = "DOCX Content"
        ws1['A1'].font = header_font
        ws1['A1'].fill = header_fill
        ws1['A1'].alignment = Alignment(horizontal='center', vertical='center')
        
        row = 2
        for line in docx_content:
            ws1[f'A{row}'] = line
            row += 1
        
        ws1.column_dimensions['A'].width = 100
        
        # Sheet 2: CSV Data
        ws2 = wb.create_sheet(title="CSV Data")
        
        for col_idx, column in enumerate(csv_df.columns, start=1):
            cell = ws2.cell(row=1, column=col_idx, value=column)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        for row_idx, row in enumerate(csv_df.itertuples(index=False), start=2):
            for col_idx, value in enumerate(row, start=1):
                ws2.cell(row=row_idx, column=col_idx, value=value)
        
        for column in ws2.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws2.column_dimensions[column_letter].width = adjusted_width
        
        # Sheet 3: Summary
        ws3 = wb.create_sheet(title="Summary")
        ws3['A1'] = "File Processing Summary"
        ws3['A1'].font = Font(bold=True, size=14)
        ws3['A3'] = "Processing Date:"
        ws3['B3'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        ws3['A4'] = "DOCX Lines:"
        ws3['B4'] = len(docx_content)
        ws3['A5'] = "CSV Rows:"
        ws3['B5'] = len(csv_df)
        ws3['A6'] = "CSV Columns:"
        ws3['B6'] = len(csv_df.columns)
        ws3.column_dimensions['A'].width = 20
        ws3.column_dimensions['B'].width = 30
        
        # Save to stream
        wb.save(output_stream)
        output_stream.seek(0)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"output_{timestamp}.xlsx"
        
        return send_file(
            output_stream,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=output_filename
        )
    
    except Exception as e:
        return jsonify({
            'error': f'Error processing files: {str(e)}'
        }), 500


    

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
